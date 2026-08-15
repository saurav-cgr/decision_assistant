from pathlib import Path
from shutil import copyfile
from zipfile import ZIP_DEFLATED, ZipFile

import pytest

from decision_assistant.ingestion.parsers import DocumentParseError, parse_document


DOCX_FIXTURE = Path("tests/fixtures/decision.docx")


def test_docx_parser_preserves_paragraph_ranges_and_linearizes_table_cells() -> None:
    parsed = parse_document(DOCX_FIXTURE)

    assert [block.text for block in parsed.blocks] == [
        "Decision Record",
        "Authentication was postponed.",
        "Owner",
        "Maya",
        "Status",
        "Proposed",
        "Review next quarter.",
    ]
    assert [block.locator for block in parsed.blocks] == [
        {"kind": "docx_paragraphs", "start": number, "end": number}
        for number in range(1, 8)
    ]
    assert all(
        parsed.content[block.start_offset : block.end_offset] == block.text
        for block in parsed.blocks
    )


def test_docx_heading_styles_map_and_table_cells_are_table_cell(
    tmp_path: Path,
) -> None:
    from docx import Document as open_docx

    path = tmp_path / "styled.docx"
    document = open_docx()
    document.add_heading("Architecture", level=1)
    document.add_paragraph("Body text")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Maya"
    document.save(path)

    parsed = parse_document(path)

    headings = [block for block in parsed.blocks if block.block_type == "heading"]
    assert [block.text for block in headings] == ["Architecture"]
    assert headings[0].attributes["level"] == 1
    assert headings[0].boundary_before == "none"
    assert headings[0].group_path == ("heading-1:architecture#1",)

    table_cells = [block for block in parsed.blocks if block.block_type == "table_cell"]
    assert {block.text for block in table_cells} == {"Owner", "Maya"}
    assert all(block.boundary_before in {"none", "soft"} for block in table_cells)


def test_docx_embedded_macro_payload_is_never_executed(tmp_path: Path) -> None:
    instrumented = tmp_path / "instrumented.docx"
    marker = tmp_path / "macro-executed"
    copyfile(DOCX_FIXTURE, instrumented)
    payload = (
        "from pathlib import Path; "
        f"Path({str(marker)!r}).write_text('executed')"
    ).encode()
    with ZipFile(instrumented, "a", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/vbaProject.bin", payload)
        archive.writestr("word/scripts/untrusted.py", payload)

    parsed = parse_document(instrumented)

    assert marker.exists() is False
    assert "macro-executed" not in parsed.content
    assert parsed.blocks[0].text == "Decision Record"


def test_corrupt_docx_returns_parser_specific_error(tmp_path: Path) -> None:
    corrupt = tmp_path / "corrupt.docx"
    corrupt.write_bytes(b"not an OOXML package")

    with pytest.raises(DocumentParseError) as error:
        parse_document(corrupt)

    assert error.value.code == "docx_parse_failed"
