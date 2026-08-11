"""Build deterministic binary documents for the fictional Atlas benchmark."""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIRECTORY = PROJECT_ROOT / "sample_data" / "atlas"
DOCX_PATH = OUTPUT_DIRECTORY / "03-auth-rollout.docx"
PDF_PATH = OUTPUT_DIRECTORY / "04-q3-planning.pdf"

INK = RGBColor(31, 41, 55)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(90, 98, 108)


def _set_run_font(
    run,
    *,
    name: str = "Arial",
    size: float = 11,
    color: RGBColor = INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 12, 6),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, RGBColor(31, 77, 120), 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _add_metadata_line(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, bold=True)
    _set_run_font(paragraph.add_run(value))


def build_docx() -> None:
    document = Document()
    _configure_docx_styles(document)
    properties = document.core_properties
    properties.title = "Atlas Authentication Rollout Decision Memo"
    properties.subject = "Fictional Atlas benchmark source"
    properties.author = "Atlas Platform Team"
    properties.created = datetime(2026, 7, 8, 9, 0, tzinfo=UTC)
    properties.modified = datetime(2026, 7, 8, 9, 0, tzinfo=UTC)

    for line in (
        "---",
        "title: Atlas Authentication Rollout Decision Memo",
        "date: 2026-07-08",
        "participants: [Priya Nair, Jonah Reed, Elena Park]",
        "source_type: decision_memo",
        "project: Atlas",
        "---",
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        _set_run_font(paragraph.add_run(line), size=8.5, color=MUTED)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(
        title.add_run("DECISION MEMO"),
        size=23,
        color=RGBColor(0, 0, 0),
        bold=True,
    )
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_run_font(
        subtitle.add_run("Atlas Authentication Rollout Revision"),
        size=14,
        color=MUTED,
    )

    _add_metadata_line(document, "Date", "July 8, 2026")
    _add_metadata_line(document, "Owner", "Priya Nair")
    _add_metadata_line(document, "Status", "Accepted revision")

    document.add_heading("Decision", level=1)
    document.add_paragraph(
        "Decision: Begin the employee-only authentication beta on July 22, 2026. "
        "Status: active. Priya Nair owns the rollout, and Jonah Reed owns security "
        "approval."
    )
    document.add_paragraph(
        "This accepted revision supersedes the June 12 proposal to begin the internal "
        "beta on July 15. The date moved by one week so all six authorization audit "
        "events can complete integration testing."
    )
    document.add_paragraph(
        "Public customer authentication remains postponed to Q4 2026. The internal "
        "beta does not supersede the May 20 public-rollout postponement."
    )

    document.add_heading("Reason and evidence", level=1)
    document.add_paragraph(
        "The authorization audit trail must record the actor, subject, previous value, "
        "new value, and timestamp for every role or permission change. Five events "
        "pass; the permission-override event still fails the replay test."
    )

    document.add_heading("Alternatives considered", level=1)
    document.add_paragraph(
        "Keep the July 15 date and disable permission overrides during beta. Status: "
        "rejected. The team would not learn whether the complete authorization path "
        "works under realistic use."
    )
    document.add_paragraph(
        "Cancel the internal beta and wait for Q4. Status: rejected. The team chose a "
        "one-week delay because limited employee evidence is useful before public "
        "release."
    )

    document.add_heading("Exit criteria", level=1)
    document.add_paragraph(
        "Security approval requires all six event tests, no critical login defects, "
        "and a documented rollback. A separate decision is required before any "
        "customer receives authentication access."
    )

    document.save(DOCX_PATH)


def _pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    return Paragraph(escaped, style)


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "AtlasBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        textColor=HexColor("#1F2937"),
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    metadata = ParagraphStyle(
        "AtlasMetadata",
        parent=body,
        fontSize=8.5,
        leading=10,
        textColor=HexColor("#5A626C"),
        spaceAfter=1,
    )
    title = ParagraphStyle(
        "AtlasTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=HexColor("#0B2545"),
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    subtitle = ParagraphStyle(
        "AtlasSubtitle",
        parent=body,
        fontName="Helvetica",
        fontSize=12,
        leading=15,
        textColor=HexColor("#5A626C"),
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    heading = ParagraphStyle(
        "AtlasHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        textColor=HexColor("#2E74B5"),
        spaceBefore=10,
        spaceAfter=6,
    )

    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Atlas Q3 Planning Readout",
        author="Atlas Delivery Team",
        subject="Fictional Atlas benchmark source",
        invariant=1,
    )
    story = []
    for line in (
        "---",
        "title: Atlas Q3 Planning Readout",
        "date: 2026-07-18",
        "participants: [Marco Silva, Elena Park, Dana Wu]",
        "source_type: planning_readout",
        "project: Atlas",
        "---",
    ):
        story.append(_pdf_paragraph(line, metadata))
    story.extend(
        [
            Spacer(1, 12),
            _pdf_paragraph("ATLAS Q3 PLANNING READOUT", title),
            _pdf_paragraph("Delivery notes - July 18, 2026", subtitle),
            _pdf_paragraph("Authentication", heading),
            _pdf_paragraph(
                "Planning note: Public authentication is active for Q3 and Marco "
                "Silva owns the rollout. Status: active.",
                body,
            ),
            _pdf_paragraph(
                "This statement conflicts with the approved July 8 decision memo, "
                "which limits access to an employee-only beta and keeps the public "
                "rollout postponed. The planning group did not record an approval or "
                "supporting security evidence for its statement.",
                body,
            ),
            _pdf_paragraph("Offline model packaging", heading),
            _pdf_paragraph(
                "Decision: Do not bundle Ollama model weights in the application "
                "images. Status: active. Dana Wu owns setup documentation.",
                body,
            ),
            _pdf_paragraph(
                "Reason: Model weights change independently and would make application "
                "images too large. Installation will use an explicit model-pull step.",
                body,
            ),
            _pdf_paragraph("Evaluation target", heading),
            _pdf_paragraph(
                "Decision: The expected source must appear in the top five retrieval "
                "results for at least 80 percent of answerable benchmark questions. "
                "Status: active. Elena Park owns acceptance.",
                body,
            ),
            _pdf_paragraph(
                "Alternative considered: Report answer quality without a retrieval "
                "threshold. Status: rejected because generation metrics cannot reveal "
                "whether the correct evidence was retrieved.",
                body,
            ),
        ]
    )
    document.build(story)


def main() -> None:
    OUTPUT_DIRECTORY.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(f"Generated {DOCX_PATH.relative_to(PROJECT_ROOT)}")
    print(f"Generated {PDF_PATH.relative_to(PROJECT_ROOT)}")


if __name__ == "__main__":
    main()
